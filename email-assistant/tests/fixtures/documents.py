"""Builders for real documents used in extraction tests.

Genuine PDF/DOCX/XLSX bytes, produced by the same libraries the wild uses —
a mocked parser would prove nothing about whether extraction works.
"""

from __future__ import annotations

import io
import zipfile


def make_pdf(pages: list[str]) -> bytes:
    """A minimal but valid PDF with a real text layer.

    Hand-built rather than pulled from a generator dependency: the format is
    simple enough, and this keeps the test suite free of another library.
    """
    objects: list[bytes] = []
    page_count = len(pages)

    kids = " ".join(f"{4 + 2 * i} 0 R" for i in range(page_count))
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode())
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    for text in pages:
        escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1", "replace")
        objects.append(
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 3 0 R >> >> /Contents "
            + str(len(objects) + 2).encode()
            + b" 0 R >>"
        )
        objects.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
        )

    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = []
    for index, body in enumerate(objects, start=1):
        offsets.append(out.tell())
        out.write(f"{index} 0 obj\n".encode() + body + b"\nendobj\n")

    xref_at = out.tell()
    out.write(f"xref\n0 {len(objects) + 1}\n".encode())
    out.write(b"0000000000 65535 f \n")
    for offset in offsets:
        out.write(f"{offset:010d} 00000 n \n".encode())
    out.write(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF".encode()
    )
    return out.getvalue()


def make_scanned_pdf() -> bytes:
    """A PDF with a page but no text layer — what a scanner produces."""
    return make_pdf([""])


def make_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table:
        added = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for col_index, value in enumerate(row):
                added.cell(row_index, col_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_xlsx(sheets: dict[str, list[list]]) -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)
    for title, rows in sheets.items():
        sheet = workbook.create_sheet(title=title)
        for row in rows:
            sheet.append(row)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def make_broken_zip() -> bytes:
    """Looks like an Office file by magic bytes, but is not one."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("random.txt", "not an office document")
    return buffer.getvalue()


def make_docx_with_revisions(
    *,
    before: str = "Zmluvna pokuta je ",
    deleted: str = "5000 EUR",
    inserted: str = "2000 EUR",
    after: str = ".",
    delete_author: str = "Advokat",
    insert_author: str = "Protistrana",
    comment: tuple[str, str] | None = None,
) -> bytes:
    """A .docx carrying real tracked changes, and optionally a comment.

    Built by injecting genuine ``w:ins`` / ``w:del`` markup into a document
    python-docx produced, because python-docx cannot author revisions itself.
    """
    import io
    import re
    import zipfile

    import docx

    document = docx.Document()
    document.add_paragraph("PLACEHOLDER")
    buffer = io.BytesIO()
    document.save(buffer)
    base = buffer.getvalue()

    def esc(value: str) -> str:
        return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    body = "<w:p>"
    if before:
        body += f"<w:r><w:t xml:space='preserve'>{esc(before)}</w:t></w:r>"
    if deleted:
        body += (
            f'<w:del w:id="101" w:author="{esc(delete_author)}" '
            f'w:date="2026-08-01T10:00:00Z">'
            f"<w:r><w:delText xml:space='preserve'>{esc(deleted)}</w:delText></w:r></w:del>"
        )
    if inserted:
        body += (
            f'<w:ins w:id="102" w:author="{esc(insert_author)}" '
            f'w:date="2026-08-02T11:00:00Z">'
            f"<w:r><w:t xml:space='preserve'>{esc(inserted)}</w:t></w:r></w:ins>"
        )
    if after:
        body += f"<w:r><w:t xml:space='preserve'>{esc(after)}</w:t></w:r>"
    body += "</w:p>"

    with zipfile.ZipFile(io.BytesIO(base)) as archive:
        document_xml = archive.read("word/document.xml").decode()
        others = {
            name: archive.read(name) for name in archive.namelist() if name != "word/document.xml"
        }

    document_xml = re.sub(r"<w:p\b.*?</w:p>", body, document_xml, count=1, flags=re.DOTALL)

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", document_xml)
        for name, payload in others.items():
            archive.writestr(name, payload)
        if comment:
            author, text = comment
            archive.writestr(
                "word/comments.xml",
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/'
                'wordprocessingml/2006/main">'
                f'<w:comment w:id="1" w:author="{esc(author)}" '
                f'w:date="2026-08-03T09:00:00Z">'
                f"<w:p><w:r><w:t>{esc(text)}</w:t></w:r></w:p>"
                "</w:comment></w:comments>",
            )
    return out.getvalue()


def make_zip(members: dict[str, bytes], compress: bool = True) -> bytes:
    """A ZIP holding *members*, in the given order."""
    import zipfile

    buffer = io.BytesIO()
    mode = zipfile.ZIP_DEFLATED if compress else zipfile.ZIP_STORED
    with zipfile.ZipFile(buffer, "w", mode) as archive:
        for name, payload in members.items():
            archive.writestr(name, payload)
    return buffer.getvalue()


def make_asice(payloads: dict[str, bytes]) -> bytes:
    """A container shaped like the real ASiC-E samples in this repository.

    mimetype first, payloads at the root, signatures and manifest under
    META-INF — the layout verified against sample1.asice and sample2.asice.
    """
    members: dict[str, bytes] = {"mimetype": b"application/vnd.etsi.asic-e+zip"}
    members.update(payloads)
    entries = "".join(
        f'<manifest:file-entry manifest:full-path="{name}" manifest:media-type="application/pdf"/>'
        for name in payloads
    )
    members["META-INF/manifest.xml"] = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">'
        '<manifest:file-entry manifest:full-path="/" '
        'manifest:media-type="application/vnd.etsi.asic-e+zip"/>'
        f"{entries}</manifest:manifest>"
    ).encode()
    members["META-INF/signatures0.xml"] = (
        b'<?xml version="1.0" encoding="UTF-8"?>'
        b'<XAdESSignatures xmlns="http://uri.etsi.org/2918/v1.2.1#">'
        b'<Signature xmlns="http://www.w3.org/2000/09/xmldsig#">'
        b"<SignatureValue>PODPIS-NEOVERENY</SignatureValue>"
        b"</Signature></XAdESSignatures>"
    )
    return make_zip(members)


def make_zip_bomb(uncompressed_bytes: int = 20 * 1024 * 1024) -> bytes:
    """A small archive whose single member declares an enormous size."""
    return make_zip({"bomb.txt": b"\0" * uncompressed_bytes})


def make_ics(body: str, newline: str = "\r\n", encoding: str = "utf-8") -> bytes:
    """A VCALENDAR wrapper around *body*, with the given line ending."""
    text = f"BEGIN:VCALENDAR\nVERSION:2.0\nPRODID:-//Test//EN\n{body}\nEND:VCALENDAR\n"
    return text.replace("\n", newline).encode(encoding)


def make_isdoc(number: str = "2026045", total: str = "1250.50") -> bytes:
    """An ISDOC-shaped invoice: the real namespace, a plausible element tree."""
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Invoice xmlns="http://isdoc.cz/namespace/2013" version="6.0.2">'
        f"<ID>{number}</ID>"
        "<IssueDate>2026-03-01</IssueDate>"
        "<AccountingSupplierParty><Party><PartyName>"
        "<Name>Stavby s.r.o.</Name></PartyName></Party></AccountingSupplierParty>"
        "<PaymentMeans><Payment><Details>"
        "<IBAN>SK3112000000198742637541</IBAN>"
        "<VariableSymbol>2026045</VariableSymbol>"
        "</Details></Payment></PaymentMeans>"
        f"<LegalMonetaryTotal><PayableAmount>{total}</PayableAmount></LegalMonetaryTotal>"
        "</Invoice>"
    ).encode()


def make_png(width: int, height: int) -> bytes:
    """A PNG header with real dimensions — enough for a size probe."""
    import struct
    import zlib

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    chunk = struct.pack(">I", len(ihdr)) + b"IHDR" + ihdr
    chunk += struct.pack(">I", zlib.crc32(b"IHDR" + ihdr) & 0xFFFFFFFF)
    return b"\x89PNG\r\n\x1a\n" + chunk


def make_gif(width: int, height: int) -> bytes:
    import struct

    return b"GIF89a" + struct.pack("<HH", width, height) + b"\x00\x00\x00"


def make_jpeg(width: int, height: int) -> bytes:
    """A JPEG whose APP0 and SOF0 segments are both well formed.

    The declared segment length has to match the payload: a reader walks the
    chain by it, and one wrong length skips the frame header entirely.
    """
    import struct

    # APP0/JFIF: identifier(5) version(2) units(1) density(4) thumbnail(2).
    app0_payload = b"JFIF\x00" + bytes([1, 1, 0]) + struct.pack(">HH", 72, 72) + b"\x00\x00"
    app0 = b"\xff\xe0" + struct.pack(">H", 2 + len(app0_payload)) + app0_payload

    # SOF0: precision(1) height(2) width(2) components(1) then 3 bytes each.
    sof_payload = struct.pack(">BHHB", 8, height, width, 1) + bytes([1, 0x11, 0])
    sof = b"\xff\xc0" + struct.pack(">H", 2 + len(sof_payload)) + sof_payload

    return b"\xff\xd8" + app0 + sof + b"\xff\xd9"


def make_locked_pdf(password: str = "s3cret") -> bytes:
    """A PDF whose password we do not have."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.encrypt(password)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()
